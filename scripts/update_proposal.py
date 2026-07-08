"""
GURO Proposal Document Updater
Updates docs/03-IT140_*.docx to reflect the actual Laravel+MySQL+Gemini stack,
fixes all Firebase/NoSQL references, adds missing architectures and terms.
"""

import copy
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'docs/03-IT140_CapstoneProjectAndResearch1_Proposal-Final-Document-Temaplate_04-18-2026.docx'
DST = 'docs/GURO_Proposal_Updated.docx'

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

def get_text(para):
    return ''.join(t.text for t in para._p.iter(f'{{{W}}}t') if t.text)

def replace_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run carrying new_text,
    preserving paragraph-level style."""
    p = para._p
    for child in list(p):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink', 'ins', 'del', 'bookmarkStart', 'bookmarkEnd'):
            # keep bookmark elements? actually remove runs only
            if tag in ('r', 'hyperlink', 'ins', 'del'):
                p.remove(child)
    r = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = new_text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_el)
    p.append(r)

def make_para_elem(text, style_val=None, bold=False):
    """Create a fresh w:p XML element."""
    new_p = OxmlElement('w:p')
    if style_val:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_val)
        pPr.append(pStyle)
        new_p.append(pPr)
    r = OxmlElement('w:r')
    if bold:
        rPr = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr.append(b)
        r.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t_el)
    new_p.append(r)
    return new_p

def insert_paras_after(ref_p_elem, paras_data):
    """Insert list of (text, style_val) after ref_p_elem (an lxml element).
    Returns the last inserted element."""
    cur = ref_p_elem
    for text, style_val in paras_data:
        new_p = make_para_elem(text, style_val)
        cur.addnext(new_p)
        cur = new_p
    return cur

def delete_elem(elem):
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

# ──────────────────────────────────────────────────────────────
# Load document
# ──────────────────────────────────────────────────────────────
doc = Document(SRC)
body = doc.element.body

# Build a flat list of all body children (paragraphs AND tables)
def body_children():
    return list(body)

# ──────────────────────────────────────────────────────────────
# 1. Chapter 2 intro — fix "web-based application" → "cross-platform application"
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'web-based application GURO: Guided Unified Remote Online' in t:
        new_t = t.replace(
            'web-based application GURO: Guided Unified Remote Online - A Rule-Based Adaptive Tutoring Web Application',
            'cross-platform application GURO: Guided Unified Remote Online – A Rule-Based Adaptive Tutoring Cross-Platform Application'
        )
        replace_para_text(para, new_t)
        print('[FIX 1] Ch2 intro web-based -> cross-platform done')
        break

# ──────────────────────────────────────────────────────────────
# 2. Chapter 1.3 — "five key technical architectures" → "seven"
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'five key technical architectures' in t:
        replace_para_text(para, t.replace('five key technical architectures', 'seven key technical architectures'))
        print('[FIX 2] Five -> seven architectures done')
        break

# ──────────────────────────────────────────────────────────────
# 3. After "Offline-First Mobile Platform" paragraph — insert 2 new architectures
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if t.startswith('Offline-First Mobile Platform'):
        new_paras = [
            (
                'AI-Powered Content Generation – An integration with the Google Gemini 2.5 Flash generative AI API that accepts teacher-supplied lesson text or a PDF document and returns a structured JSON payload containing a child-friendly study guide (introduction, vocabulary definitions, and key takeaway summary) and a full set of tiered diagnostic questions (Easy, Average, and Difficult) with bilingual answer feedback. This component offloads content authoring from teachers, enabling rapid creation of curriculum-aligned assessment material without manual question writing.',
                None
            ),
            (
                'Teacher and Parent Web Portal – A React + Vite web application providing role-specific dashboards for teachers (classroom analytics, classroom setup and invite-code generation, manual lesson builder, and AI lesson ingestion pipeline), parents (student progress monitoring via a six-digit access code, activity heatmap, badge case, and tutor report), and system administrators (system-wide overview dashboard and lesson ingestor). The portal communicates with the same Laravel REST API backend as the mobile application and supports dark/light theme toggling and a collapsible sidebar.',
                None
            ),
        ]
        insert_paras_after(para._p, new_paras)
        print('[FIX 3] New architectures inserted after Offline-First paragraph')
        break

# ──────────────────────────────────────────────────────────────
# 4. Specific Objective #2 — fix "client-side, rule based adaptive tutoring engine"
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if t.startswith('Develop a client-side, rule based adaptive tutoring engine which adapt the difficulty level'):
        replace_para_text(para,
            'Develop a server-assisted adaptive learning pipeline that combines: (a) a client-side rule-based difficulty-routing engine embedded in the mobile application that dynamically adjusts the difficulty tier and sequence of assessment items based on the student’s real-time quiz responses; and (b) a server-side AI content generation component powered by the Google Gemini 2.5 Flash API that produces structured study content and tiered diagnostic questions from teacher-uploaded lesson text or PDF materials, thereby personalizing the learning path for each individual student. [6, 7]'
        )
        print('[FIX 4] SO #2 client-side fix done')
        break

# ──────────────────────────────────────────────────────────────
# 5. Scope and Limitation — update to mention web portal + AI dependency
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if t.startswith('The scope of this project is on the design and development of the Guided Unified Remote Online (GURO)'):
        replace_para_text(para,
            'The scope of this project is on the design and development of the Guided Unified Remote Online (GURO) cross-platform application, which is a rule-based adaptive tutoring system targeting Grade 4–6 elementary students in Mathematics and English. The system encompasses three client layers: (1) an offline-first React Native mobile application for student-facing learning, assessment, and progress tracking; (2) a React + Vite web portal providing teacher analytics, classroom management, lesson ingestion, and parent monitoring dashboards; and (3) a Laravel 11 REST API backend with a MySQL relational database that handles authentication, classroom management, AI-powered content generation via Google Gemini 2.5 Flash, item bank storage, and progress telemetry synchronization. The application supports adaptive tutoring, outcome-driven assessments, predictive performance analysis, and randomized content delivery. Implementation includes a client-side rule-based adaptive engine, an AI-assisted server-side content generation pipeline, a classroom invite-code pairing system, offline-first SQLite local storage with background sync, a gamification layer (XP points, virtual stars, streak counters, collectible badges, and a mascot outfit shop), and a parental control module. System evaluation will include structured system testing, UAT with Grade 4–6 teachers and students, and performance benchmarking on low-cost Android devices. [1, 4, 6]'
        )
        print('[FIX 5a] Scope paragraph updated')
        break

for para in doc.paragraphs:
    t = get_text(para)
    if 'The GURO application also features no additional features beyond a mobile platform' in t:
        replace_para_text(para,
            'Additional limitations of this study include the following. The Teacher Portal and Admin Console web application require internet connectivity and are not offline-capable; this constraint is intentional since teacher and admin workflows depend on real-time data synchronization. The AI content generation component depends on a valid Google Gemini API key and active internet access at the server level; content generation will be unavailable if the API service is unreachable. System evaluation, usability testing, and minimal performance benchmarking via controlled pilot with target users constitute the primary evaluation approach, rather than long-term longitudinal deployment or nationwide classroom rollout. Despite being designed to assist students in GIDA communities, the mobile application still requires the availability of a shared or personal Android or iOS device; device access limitations among target users are therefore outside the scope of this study. These constraints define the technological and operational boundaries of the research and ensure the scope remains achievable within the IT 140 capstone timeline.'
        )
        print('[FIX 5b] Limitations paragraph updated')
        break

# ──────────────────────────────────────────────────────────────
# 6. Definition of Terms — add 5 new entries after last existing entry
# ──────────────────────────────────────────────────────────────
# The DoT table ends with "Diagnostic assessment" definition. After the table there is a blank paragraph.
# We find the paragraph AFTER the DoT table by looking for the text of the last definition.
# Since the table is a w:tbl, we work at the body-children level.

# Strategy: find the w:tbl that contains "Diagnostic assessment", then insert after it.
for child in body_children():
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'tbl':
        # Check if this table contains "Diagnostic assessment"
        text_in_tbl = ''.join(t.text for t in child.iter(f'{{{W}}}t') if t.text)
        if 'Diagnostic assessment' in text_in_tbl and 'Rule-Based Adaptive' in text_in_tbl:
            # This is the Definition of Terms table. Insert new entries after it.
            new_terms = [
                ('', None),  # blank separator
                ('The following additional terms are introduced in this study to reflect the current system implementation:', None),
                ('', None),
                ('Gemini API – The Google Gemini 2.5 Flash generative AI API used in GURO’s server-side content generation pipeline. It accepts structured prompts containing lesson text or PDF content and returns a structured JSON response conforming to a predefined schema that includes a child-friendly study guide and a set of diagnostic questions organized by difficulty tier.', None),
                ('', None),
                ('Classroom Invite Code – A formatted alphanumeric identifier (e.g., MAT-G5-XYZ) generated by the teacher through the GURO Teacher Portal when creating a classroom session. The code encodes the subject prefix, grade level, and a random three-character suffix. Students enter this code in the mobile application to pair their device to the teacher’s classroom and receive the classroom’s custom item bank.', None),
                ('', None),
                ('Item Bank – The structured JSON data store of all curriculum-aligned questions and study content used by the GURO application. Items are indexed hierarchically by subject, grade level, and topic, and each topic node contains a studyContent object and questions organized by difficulty tier (Easy, Average, Difficult) and category (Multiple-Choice, Paragraph Comprehension, Figures of Speech). A global item bank is maintained as a JSON file on the server; classroom-specific custom banks are stored as JSON columns in the MySQL classrooms table.', None),
                ('', None),
                ('Parent Access Code – A deterministic six-digit numeric code derived from the student’s device identifier using a salted positional hash function. The code allows a parent or guardian to access a read-only view of the student’s progress data through the GURO Parent Web Portal without requiring a registered user account.', None),
                ('', None),
                ('Guest-to-Registered Account Promotion – A workflow in the GURO mobile application that allows a student who has been using the app as an anonymous guest (identified by a local device ID) to register a cloud account with an email address and password. Upon promotion, the system migrates all previously recorded progress events from the guest device ID to the new registered student ID, preserving the learner’s full learning history.', None),
            ]
            insert_paras_after(child, new_terms)
            print('[FIX 6] New DoT entries inserted after table')
            break

# ──────────────────────────────────────────────────────────────
# 7. Chapter 3 — Fix Backend Platform (Firebase → Laravel 11)
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'Backend Platform' in t and 'Firebase (Authentication, Firestore, Cloud Functions, Cloud Storage)' in t:
        replace_para_text(para, 'Backend Platform – Laravel 11 (PHP) and MySQL')
        print('[FIX 7a] Ch3 Backend Platform header fixed')
        break

for para in doc.paragraphs:
    t = get_text(para)
    if 'Firebase Firestore has been selected as the central data repository' in t:
        replace_para_text(para,
            'Laravel 11 was selected as the backend REST API framework due to its built-in support for database migrations, Eloquent ORM, HTTP client integration (used for the Gemini API calls), and Laravel Sanctum for token-based authentication. MySQL was chosen as the relational database engine for its reliability, widespread deployment support, and compatibility with Laravel’s migration toolchain. The relational model was preferred over a NoSQL alternative because the GURO data model — users, classrooms, and progress logs — has well-defined, structured relationships and benefits from referential integrity constraints, indexed queries, and SQL-based aggregation for analytics. Studies comparing relational and NoSQL databases for educational data applications indicate that relational databases provide superior query consistency and data integrity for structured learner performance records [5].'
        )
        print('[FIX 7b] Firebase Firestore description replaced with Laravel/MySQL')
        break

for para in doc.paragraphs:
    t = get_text(para)
    if 'This enables the client to act on a local cache and resolve discrepancies once network connectivity is restored' in t:
        replace_para_text(para,
            'Laravel Sanctum issues short-lived bearer tokens upon login or registration. All protected API routes (lesson generation, classroom management, item bank saving) require a valid Sanctum token in the Authorization header. Student progress synchronization and public classroom verification routes are intentionally unauthenticated to allow offline-first student devices to sync without a prior login step. PBKDF2-SHA512 with a random 16-byte salt and 1,000 iterations is used for password hashing, maintaining compatibility with the prior Express.js hashing scheme used during early development.'
        )
        print('[FIX 7c] Firebase client cache paragraph replaced with Sanctum description')
        break

# ──────────────────────────────────────────────────────────────
# 8. Chapter 3 — Fix Offline Storage (IndexedDB → expo-sqlite)
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'Offline Storage' in t and 'IndexedDB, AsyncStorage, and SQLite' in t:
        replace_para_text(para, 'Offline Storage – expo-sqlite and AsyncStorage')
        print('[FIX 8a] Offline Storage header fixed')
        break

for para in doc.paragraphs:
    t = get_text(para)
    if 'Empirical research comparing applications utilizing' in t and 'offline-first architecture' in t:
        replace_para_text(para,
            'Offline-first architecture research confirms that applications that persist data locally before syncing to a remote server provide significantly better user experience in low-connectivity environments [4]. GURO’s mobile application uses expo-sqlite (a SQLite wrapper for Expo-managed React Native projects) as its primary local data store for three categories of data: student progress events (quiz scores, topics attempted, timestamps), item bank cache (the JSON question library indexed by subject/grade/topic), and parental settings (PIN, daily time limits, unlocked badges, XP points, virtual stars, mascot outfit, voice guide theme). Zustand with the zustand/middleware persist adapter and React Native AsyncStorage handles in-memory state management and lightweight key-value persistence (session tokens, theme preference, sidebar state). All progress events written to SQLite are queued for background synchronization to the Laravel API via the syncProgressNow action; events are marked as synced locally once the server acknowledges receipt, preventing duplicate upload on subsequent sync attempts.'
        )
        print('[FIX 8b] Offline Storage description updated')
        break

# ──────────────────────────────────────────────────────────────
# 9. Chapter 3 — Insert Gemini API section after Rule-Based Adaptive Engine section
# ──────────────────────────────────────────────────────────────
gemini_inserted = False
for para in doc.paragraphs:
    t = get_text(para)
    if t.startswith('Gamification Elements'):
        # Insert Gemini section BEFORE Gamification (i.e., after the preceding paragraph)
        gemini_paras = [
            ('AI Content Generation – Google Gemini 2.5 Flash', None),
            (
                'Google Gemini 2.5 Flash was integrated as the AI content generation backbone of the GURO lesson ingestion pipeline. When a teacher submits lesson material through the Teacher Portal — either as pasted lesson text or a base64-encoded PDF upload — the Laravel GeminiService constructs a structured prompt and sends it to the Gemini API v1beta endpoint with a JSON response schema enforcing the exact output structure required by the application. The schema mandates a studyContent object (containing an age-appropriate introduction, an array of vocabulary term definitions with examples, and a summary bullet list) and a questions array (containing Multiple-Choice, Paragraph Comprehension, and Figures of Speech items stratified across Easy, Average, and Difficult difficulty tiers, each with four options, a correct answer, and bilingual English/Filipino feedback). The use of Gemini’s responseMimeType: application/json with a responseSchema eliminates the need for post-processing or regex extraction, delivering immediately parse-ready structured data. A generation temperature of 0.2 is used to minimize hallucinations while still producing varied question formulations. This approach significantly accelerates teacher-side content creation, reducing lesson preparation time from hours to minutes.',
                None
            ),
        ]
        # Find the paragraph right before this one and insert after it
        prev_para = None
        for p2 in doc.paragraphs:
            if p2._p is para._p:
                break
            prev_para = p2
        if prev_para is not None:
            insert_paras_after(prev_para._p, gemini_paras)
            gemini_inserted = True
            print('[FIX 9] Gemini API section inserted')
        break

# ──────────────────────────────────────────────────────────────
# 10. Chapter 4 System Architecture — fix Data Layer (IndexedDB/Firebase refs)
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'Data Layer' in t and 'IndexedDB within PWAs' in t:
        replace_para_text(para,
            'Data Layer is divided between on-device data storage and cloud data storage. On-device storage uses expo-sqlite within the Student Mobile App (for progress events, item bank cache, parental settings, and achievement data) and AsyncStorage (for Zustand state persistence and lightweight key-value data). Cloud data storage uses a MySQL relational database managed through the Laravel 11 backend, comprising three core tables: users (authentication credentials and role information), classrooms (teacher-created sessions with subject, grade level, invite code, expiry, and custom item bank as a JSON column), and progress_logs (individual student quiz attempt records including score, topic, grade level, and sync timestamp). All client-to-cloud data transfer is mediated through the GURO REST API over HTTPS.'
        )
        print('[FIX 10] Ch4 Data Layer fixed')
        break

# ──────────────────────────────────────────────────────────────
# 11. Chapter 4 Security — fix Firebase Auth → Laravel Sanctum
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'All three roles are authenticated using Firebase Authentication' in t:
        replace_para_text(para,
            'All user roles (student, teacher, parent, admin, lesson-builder, developer) are authenticated through the Laravel Sanctum token-based authentication system. Upon successful login or registration, the API issues a personal access token (stored via the personal_access_tokens table) that must be included as a Bearer token in the Authorization header of all protected API requests. Passwords are hashed using PBKDF2-SHA512 with a randomly generated 16-byte salt and 1,000 hash iterations; the stored value encodes both salt and hash in a colon-delimited format. The parent monitoring access code is a deterministic six-digit integer derived from the student’s identifier using a salted positional sum hash, allowing read-only progress access without requiring a registered account. Classroom invite codes are time-limited; the expires_at timestamp is set by the teacher at classroom creation time, and the verifyCode endpoint returns HTTP 403 once the expiry has passed, preventing unauthorized late joins.'
        )
        print('[FIX 11] Ch4 Security updated')
        break

# ──────────────────────────────────────────────────────────────
# 12. Chapter 4 Database Design — fix NoSQL/Firebase → MySQL relational
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'NoSQL database (Firebase Firestore)' in t or ('NoSQL database' in t and 'Firebase' in t):
        replace_para_text(para,
            'A MySQL relational database managed through the Laravel Eloquent ORM and schema migration system was chosen as the GURO data store. A relational model was preferred over a NoSQL alternative because the core data entities — users, classrooms, and progress logs — have clearly defined, stable schemas with enumerable attributes and benefit from SQL-based aggregation for teacher analytics queries (e.g., average score per topic per classroom). The custom_item_bank column on the classrooms table is stored as a JSON column type, taking advantage of MySQL’s native JSON support for the hierarchically structured item bank data (indexed by subject → grade → topic → difficulty → category) without requiring a fully separate document store. Database migrations are version-controlled within the Laravel project; two post-initial migrations added the parent_access_token column to the users table and the teacher_user_id foreign key to the classrooms table to support the account promotion and classroom ownership features introduced during development.'
        )
        print('[FIX 12] Ch4 Database Design updated')
        break

# ──────────────────────────────────────────────────────────────
# 13. Chapter 4 Database Fields — replace Firebase collections with MySQL tables
# ──────────────────────────────────────────────────────────────
# Find "Collection: users" paragraph and the w:tbl that follows it,
# and subsequent collection paragraphs + tables. Delete them all, insert new content.

# Step 1: Find "Database Fields" heading paragraph index in body children
# Step 2: Find "Tools and Technology" heading paragraph
# Step 3: Delete everything between them
# Step 4: Insert new MySQL schema descriptions

db_fields_heading_p = None
tools_heading_p = None

for para in doc.paragraphs:
    t = get_text(para)
    if t.strip() == 'Database Fields':
        # Check it has style heading
        db_fields_heading_p = para._p
    if t.strip() == 'Tools and Technology' and db_fields_heading_p is not None:
        tools_heading_p = para._p
        break

if db_fields_heading_p is not None and tools_heading_p is not None:
    # Collect all elements between these two paragraphs (exclusive)
    children = list(body)
    start_idx = children.index(db_fields_heading_p)
    end_idx = children.index(tools_heading_p)

    # Delete all elements between (exclusive of both endpoints)
    to_delete = children[start_idx + 1 : end_idx]
    for elem in to_delete:
        body.remove(elem)

    # Now insert new MySQL schema content after db_fields_heading_p
    mysql_content = [
        ('The following tables define the MySQL relational database schema implemented in the GURO Laravel 11 backend. The schema is managed through Laravel migrations and reflects the production state as of the current sprint.', None),
        ('', None),
        ('Table: users', None),
        ('Stores credentials and profile data for all registered user roles (student, teacher, parent, admin, lesson-builder, developer).', None),
        ('  user_id  |  VARCHAR(50)  |  PRIMARY KEY  |  NOT NULL  —  System-generated unique identifier in the format USR-XXXXXXX.', None),
        ('  email  |  VARCHAR(255)  |  UNIQUE INDEX  |  NOT NULL  —  Registered email address (lowercased, trimmed on write).', None),
        ('  password_hash  |  VARCHAR(255)  |  —  |  NOT NULL  —  PBKDF2-SHA512 hash in salt:hash colon-delimited format.', None),
        ('  name  |  VARCHAR(100)  |  —  |  NOT NULL  —  Display name of the user.', None),
        ('  role  |  VARCHAR(50)  |  —  |  NOT NULL  —  Enum: student, teacher, parent, admin, lesson-builder, developer.', None),
        ('  classroom_id  |  VARCHAR(50)  |  —  |  NULLABLE  —  Classroom the user is currently linked to.', None),
        ('  parent_access_token  |  VARCHAR(100)  |  —  |  NULLABLE  —  32-character random token generated upon student account promotion.', None),
        ('  created_at, updated_at  |  TIMESTAMP  |  —  |  Auto-managed by Eloquent.', None),
        ('', None),
        ('Table: classrooms', None),
        ('Stores teacher-created classroom sessions with subject, grade level, invite code, item bank, and expiry metadata.', None),
        ('  id  |  BIGINT  |  PRIMARY KEY (auto-increment)  |  NOT NULL.', None),
        ('  classroom_id  |  VARCHAR(50)  |  UNIQUE INDEX  |  NOT NULL  —  Formatted invite code, e.g., MAT-G5-XYZ.', None),
        ('  teacher_user_id  |  BIGINT  |  FOREIGN KEY → users.id  |  NOT NULL  —  References the owning teacher account.', None),
        ('  teacher_name  |  VARCHAR(100)  |  —  |  NOT NULL  —  Display name of the teacher.', None),
        ('  subject  |  VARCHAR(100)  |  —  |  NOT NULL  —  Subject focus (e.g., Mathematics, English).', None),
        ('  grade_level  |  INTEGER  |  —  |  NOT NULL  —  Grade level (4, 5, or 6).', None),
        ('  custom_item_bank  |  JSON  |  —  |  NULLABLE  —  Hierarchical JSON store: subject → grade → topic → difficulty → category → question array.', None),
        ('  expires_at  |  TIMESTAMP  |  —  |  NULLABLE  —  Invite code expiry; NULL means no expiry.', None),
        ('  created_at, updated_at  |  TIMESTAMP  |  —  |  Auto-managed by Eloquent.', None),
        ('', None),
        ('Table: progress_logs', None),
        ('Stores individual student quiz attempt records synced from the mobile application. Each row represents one completed quiz session for a given topic.', None),
        ('  id  |  BIGINT  |  PRIMARY KEY (auto-increment)  |  NOT NULL.', None),
        ('  event_id  |  VARCHAR(50)  |  UNIQUE INDEX  |  NOT NULL  —  Client-generated event identifier (e.g., EVT-XXXXXXX) used for deduplication on sync.', None),
        ('  student_id  |  VARCHAR(100)  |  INDEX  |  NOT NULL  —  Student device identifier or registered name slug (e.g., JUAN-DELA-CRUZ).', None),
        ('  classroom_id  |  VARCHAR(50)  |  INDEX  |  NULLABLE  —  Classroom the student was joined to at the time of the attempt.', None),
        ('  subject  |  VARCHAR(100)  |  —  |  NOT NULL  —  Subject (Mathematics or English).', None),
        ('  grade_level  |  INTEGER  |  —  |  NOT NULL  —  Grade level at the time of the attempt.', None),
        ('  topic  |  VARCHAR(255)  |  —  |  NOT NULL  —  Topic name as defined in the item bank.', None),
        ('  score  |  INTEGER  |  —  |  NOT NULL  —  Number of correct answers in the session.', None),
        ('  total_questions  |  INTEGER  |  —  |  NOT NULL  —  Total questions attempted in the session.', None),
        ('  timestamp  |  TIMESTAMP  |  —  |  NOT NULL  —  Client-side timestamp of the quiz attempt.', None),
        ('  synced_at  |  TIMESTAMP  |  —  |  DEFAULT CURRENT_TIMESTAMP  —  Server-side timestamp when the record was inserted.', None),
        ('  created_at, updated_at  |  TIMESTAMP  |  —  |  Auto-managed by Eloquent.', None),
        ('', None),
        ('Table: personal_access_tokens', None),
        ('Managed automatically by Laravel Sanctum. Stores bearer tokens issued on login/registration for all authenticated user roles.', None),
        ('  id  |  BIGINT  |  PRIMARY KEY  |  NOT NULL.', None),
        ('  tokenable_type  |  VARCHAR(255)  |  —  |  NOT NULL  —  Polymorphic model class (App\\Models\\User).', None),
        ('  tokenable_id  |  BIGINT  |  INDEX  |  NOT NULL  —  References the user’s primary key.', None),
        ('  name  |  VARCHAR(255)  |  —  |  NOT NULL  —  Token label (e.g., “app”).', None),
        ('  token  |  VARCHAR(64)  |  UNIQUE  |  NOT NULL  —  SHA256 hash of the issued token.', None),
        ('  abilities  |  TEXT  |  —  |  NULLABLE  —  JSON array of granted abilities.', None),
        ('  last_used_at, expires_at  |  TIMESTAMP  |  —  |  NULLABLE.', None),
        ('', None),
    ]
    insert_paras_after(db_fields_heading_p, mysql_content)
    print('[FIX 13] Database Fields section replaced with MySQL schema')

# ──────────────────────────────────────────────────────────────
# 14. Chapter 4 Tools and Technology — fix Firebase entry
# ──────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    t = get_text(para)
    if 'Backend Platform' in t and 'Firebase was used as the backend-as-a-service' in t:
        replace_para_text(para,
            'Backend Platform – Laravel 11 (PHP) was used as the REST API framework, providing database migrations via Eloquent ORM, Laravel Sanctum for bearer-token authentication, an HTTP client for outbound Gemini API requests, and a caching layer (1-hour TTL) for item bank reads. MySQL was used as the relational database engine for users, classrooms, and progress_logs tables.'
        )
        print('[FIX 14a] Tools Firebase Backend entry fixed')
        break

for para in doc.paragraphs:
    t = get_text(para)
    if 'Offline Storage' in t and 'IndexedDB was used for offline data caching within the PWA contexts' in t:
        replace_para_text(para,
            'Offline Storage – expo-sqlite was used as the primary on-device SQLite store within the React Native mobile application for progress events, item bank cache, parental settings, and gamification state. Zustand with the zustand/middleware persist adapter and AsyncStorage was used for lightweight in-memory state management and key-value persistence across app restarts.'
        )
        print('[FIX 14b] Tools Offline Storage entry fixed')
        break

# Also add Gemini entry near the tools list — find "Data Visualization" entry and insert before
for para in doc.paragraphs:
    t = get_text(para)
    if t.startswith('Data Visualization') and 'Chart.js' in t:
        # Insert Gemini tool entry before this
        gemini_tool_entry = make_para_elem(
            'AI Content Generation – Google Gemini 2.5 Flash (via the Gemini API v1beta) was used to power the lesson ingestion pipeline. It accepts structured prompts with lesson text or PDF content and returns JSON-schema-validated study content and diagnostic questions, eliminating manual question authoring for teachers.',
            '36'
        )
        para._p.addnext(gemini_tool_entry)
        # addnext puts it after; we need it before. Swap:
        # Actually addnext on para._p means it goes after para, but we want before it.
        # Let's remove and re-insert before:
        para._p.getparent().remove(gemini_tool_entry)
        para._p.addprevious(gemini_tool_entry)
        print('[FIX 14c] Gemini tool entry added to Tools section')
        break

# ──────────────────────────────────────────────────────────────
# 15. Chapter 4 Functional Requirements — add missing capabilities
# ──────────────────────────────────────────────────────────────
# Find last functional requirement entry and add after it
last_fr_para = None
for para in doc.paragraphs:
    t = get_text(para)
    if 'User and Device Fleet Management' in t and 'administrative management of teacher/student accounts' in t:
        last_fr_para = para
        break

if last_fr_para:
    new_fr_paras = [
        (
            'AI-Powered Lesson Generation – teachers can submit lesson text (pasted) or a PDF file (uploaded as base64) through the Lesson Ingestor, triggering an AI generation call to the Google Gemini 2.5 Flash API which returns a structured study guide and a set of difficulty-tiered diagnostic questions ready for review and saving.',
            '36'
        ),
        (
            'Classroom Item Bank Operations – teachers can selectively claim topics from the global item bank into their classroom’s custom item bank (POST /api/classroom/claim), update individual lesson content (POST /api/classroom/update-lesson), and delete specific topic entries from their classroom bank (POST /api/classroom/delete-lesson).',
            '36'
        ),
        (
            'Guest-to-Registered Account Promotion – students who begin using the mobile application as an anonymous guest can register a cloud account (POST /api/auth/promote), which migrates all previously recorded progress events from the guest device ID to the new registered student ID.',
            '36'
        ),
        (
            'Parent Progress Monitoring via Access Code – parents can view a child’s progress history through the GURO Parent Web Portal by entering a six-digit access code derived from the student’s device identifier, without requiring a registered user account.',
            '36'
        ),
        (
            'Gamification and Engagement Layer – the mobile application awards XP points (+10 per correct answer, +50 perfect quiz bonus) and virtual stars (+2 per correct answer, +10 perfect quiz bonus) per quiz session; tracks consecutive daily learning streaks; unlocks milestone badges (first_step, perfect_score, math_wizard, english_champion, streak_starter, streak_master); and provides a mascot outfit shop where students can spend virtual stars on collectible outfits.',
            '36'
        ),
        (
            'Parental Controls – parents can configure a daily screen time limit (in minutes), enforce a mathematics-before-English subject order, enable forced bilingual mode, and set a priority topic override for the student’s learning session through a PIN-protected parental controls panel in the mobile application.',
            '36'
        ),
    ]
    insert_paras_after(last_fr_para._p, new_fr_paras)
    print('[FIX 15] New functional requirements added')

# ──────────────────────────────────────────────────────────────
# Save
# ──────────────────────────────────────────────────────────────
doc.save(DST)
print(f'\n✓ Saved updated document to: {DST}')
